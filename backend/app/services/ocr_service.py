"""
ocr_service.py — Service d'extraction de texte (OCR)
=====================================================
Deux niveaux de sortie :
  - extract_text()       → texte lisible, propre (pour affichage)
  - extract_normalized() → texte normalisé (pour comparaison / vérification d'authenticité)

La normalisation garantit que le même document en PDF et en photo JPEG
produit une chaîne identique pour la vérification d'authenticité.

Exemple :
  PDF   → "BONJOUR V A ?"   → normalize → "bonjourva"
  Photo → "BONJOUR VA?"     → normalize → "bonjourva"
  → identiques ✅
"""

import io
import re
import logging
import unicodedata

import pytesseract
from PIL import Image, ImageFilter, ImageEnhance
from pdf2image import convert_from_bytes
from docx import Document

logger = logging.getLogger(__name__)

# ─── Extensions supportées ───────────────────────────────────────────────────
IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'webp', 'bmp', 'tiff', 'tif'}
PDF_EXTENSIONS   = {'pdf'}
DOCX_EXTENSIONS  = {'docx'}
ALL_EXTENSIONS   = IMAGE_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS

TESSERACT_CONFIG = r'--oem 3 --psm 3'
TESSERACT_LANG   = 'fra+eng'


class OCRService:

    # ─── Points d'entrée publics ──────────────────────────────────────────

    @staticmethod
    def extract_text(file_bytes: bytes, filename: str) -> str:
        """Texte propre et lisible — pour affichage à l'utilisateur."""
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

        if ext in IMAGE_EXTENSIONS:
            raw = OCRService._from_image(file_bytes)
        elif ext in PDF_EXTENSIONS:
            raw = OCRService._from_pdf(file_bytes)
        elif ext in DOCX_EXTENSIONS:
            raw = OCRService._from_docx(file_bytes)
        else:
            raise ValueError(
                f"Format '.{ext}' non supporté. Acceptés : {', '.join(sorted(ALL_EXTENSIONS))}"
            )

        return OCRService._clean(raw)

    @staticmethod
    def extract_normalized(file_bytes: bytes, filename: str) -> str:
        """
        Texte normalisé pour comparaison cross-format.
        Utiliser ceci pour vérifier l'authenticité d'un document.
        """
        readable = OCRService.extract_text(file_bytes, filename)
        return OCRService.normalize(readable)

    @staticmethod
    def normalize(text: str) -> str:
        """
        Normalise un texte pour comparaison d'authenticité.

        Peut être appelé sur n'importe quel texte (contenu BD ou extrait OCR).

        Exemple d'usage dans signature_service.py :
            stored   = OCRService.normalize(communique.contenu)
            uploaded = OCRService.extract_normalized(file_bytes, filename)
            is_authentic = (stored == uploaded)
        """
        if not text:
            return ''
        # 1. Minuscules
        text = text.lower()
        # 2. Décompose les accents (é→e, à→a, ç→c …)
        text = unicodedata.normalize('NFKD', text)
        text = ''.join(c for c in text if not unicodedata.combining(c))
        # 3. Garde uniquement lettres + chiffres — supprime espaces, ponctuation, sauts de ligne
        text = re.sub(r'[^a-z0-9]', '', text)
        return text

    # ─── Image ───────────────────────────────────────────────────────────

    @staticmethod
    def _from_image(file_bytes: bytes) -> str:
        try:
            image = Image.open(io.BytesIO(file_bytes))
            image = OCRService._preprocess(image)
            return pytesseract.image_to_string(image, lang=TESSERACT_LANG, config=TESSERACT_CONFIG)
        except Exception as e:
            raise RuntimeError(f"Erreur OCR image : {e}")

    @staticmethod
    def _preprocess(image: Image.Image) -> Image.Image:
        """
        Pipeline de prétraitement avancé pour photos mobiles de documents.
        Utilise OpenCV (CLAHE + débruitage + binarisation adaptative) pour
        maximiser la lisibilité OCR même sur des photos de mauvaise qualité.
        """
        import numpy as np
        try:
            import cv2
            HAS_CV2 = True
        except ImportError:
            HAS_CV2 = False

        # ── Normalise le mode couleur ─────────────────────────────────────
        if image.mode not in ('L', 'RGB', 'RGBA'):
            image = image.convert('RGB')
        if image.mode == 'RGBA':
            image = image.convert('RGB')

        # ── Upscale si l'image est trop petite (améliore beaucoup l'OCR) ──
        w, h = image.size
        if w < 1800:
            scale = 1800 / w
            image = image.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

        if not HAS_CV2:
            # Fallback PIL basique si OpenCV manquant
            image = image.convert('L')
            image = ImageEnhance.Contrast(image).enhance(2.5)
            image = image.filter(ImageFilter.SHARPEN)
            image = image.filter(ImageFilter.SHARPEN)
            return image

        # ── Conversion PIL → NumPy ─────────────────────────────────────────
        img_np = np.array(image)

        # ── Niveaux de gris ────────────────────────────────────────────────
        if len(img_np.shape) == 3:
            gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_np.copy()

        # ── Débruitage doux (conserve les bords du texte) ─────────────────
        # h=10 : force du débruitage, templateWindowSize=7, searchWindowSize=21
        denoised = cv2.fastNlMeansDenoising(gray, h=10, templateWindowSize=7, searchWindowSize=21)

        # ── CLAHE — égalisation locale du contraste ───────────────────────
        # Corrige les zones sombres/surexposées sur la même image
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        equalized = clahe.apply(denoised)

        # ── Unsharp mask (accentuation des bords) ─────────────────────────
        blurred = cv2.GaussianBlur(equalized, (0, 0), 3)
        sharpened = cv2.addWeighted(equalized, 1.5, blurred, -0.5, 0)

        # ── Binarisation adaptative de Sauvola (via Otsu + Gaussian) ──────
        # Plus robuste que le threshold global pour les documents avec ombres
        binary = cv2.adaptiveThreshold(
            sharpened,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            blockSize=31,
            C=10,
        )

        # ── Morphologie : fermeture douce pour reconnecter les lettres ────
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        cleaned = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)

        # ── Reconversion en image PIL ──────────────────────────────────────
        return Image.fromarray(cleaned)


    # @staticmethod
    # def _preprocess(image: Image.Image) -> Image.Image:
    #     from PIL import ImageOps
    #     if image.mode not in ('L', 'RGB'):
    #         image = image.convert('RGB')
    #     image = image.convert('L')
    #     # Auto-contrast (fixes dark/overexposed photos)
    #     image = ImageOps.autocontrast(image, cutoff=2)
    #     # Stronger contrast boost
    #     image = ImageEnhance.Contrast(image).enhance(2.5)
    #     # Stronger sharpening
    #     image = image.filter(ImageFilter.SHARPEN)
    #     image = image.filter(ImageFilter.SHARPEN)
    #     # Resize small images (improves OCR on low-res photos)
    #     w, h = image.size
    #     if w < 1500:
    #         image = image.resize((w * 2, h * 2), Image.LANCZOS)
    #     return image


    # ─── PDF ─────────────────────────────────────────────────────────────

    @staticmethod
    def _from_pdf(file_bytes: bytes) -> str:
        # Essai 1 : extraction directe (PDF numérique)
        direct = OCRService._pdf_direct(file_bytes)
        if len(direct.strip()) > 50:
            return direct
        # Essai 2 : OCR page par page (PDF scanné)
        return OCRService._pdf_ocr(file_bytes)

    @staticmethod
    def _pdf_direct(file_bytes: bytes) -> str:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            return '\n'.join(page.extract_text() or '' for page in reader.pages)
        except Exception:
            return ''

    @staticmethod
    def _pdf_ocr(file_bytes: bytes) -> str:
        try:
            pages = convert_from_bytes(file_bytes, dpi=300)
        except Exception as e:
            raise RuntimeError(f"Conversion PDF échouée (vérifiez poppler-utils) : {e}")
        parts = []
        for i, page in enumerate(pages, 1):
            try:
                parts.append(pytesseract.image_to_string(
                    OCRService._preprocess(page), lang=TESSERACT_LANG, config=TESSERACT_CONFIG
                ))
            except Exception as e:
                logger.warning(f"OCR page {i} : {e}")
        return '\n'.join(parts)

    # ─── DOCX ────────────────────────────────────────────────────────────

    @staticmethod
    def _from_docx(file_bytes: bytes) -> str:
        try:
            doc = Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(' | '.join(cells))
            return '\n'.join(parts)
        except Exception as e:
            raise RuntimeError(f"Erreur lecture DOCX : {e}")

    # ─── Nettoyage lisible ────────────────────────────────────────────────

    @staticmethod
    def _clean(text: str) -> str:
        """Nettoie pour affichage — garde la lisibilité humaine."""
        if not text:
            return ''
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)
        text = re.sub(r'(?<!\S)[|~^`¬§©®™](?!\S)', '', text)
        lines = [re.sub(r'  +', ' ', l).strip() for l in text.split('\n')]
        text = '\n'.join(lines)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
